import io
import re
from pathlib import Path

from docx import Document
from fastapi import HTTPException, UploadFile
from pypdf import PdfReader

from app.models.resume_extract import ResumeExtractResponse


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


async def extract_resume(file: UploadFile) -> ResumeExtractResponse:
    content = await file.read()
    extension = Path(file.filename or "").suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Supported resume formats: .txt, .pdf, .docx")

    if extension == ".txt":
        text = _extract_txt(content)
    elif extension == ".pdf":
        text = _extract_pdf(content)
    else:
        text = _extract_docx(content)

    normalized_text = _normalize_text(text)
    return ResumeExtractResponse(
        fileName=file.filename or "resume",
        contentType=file.content_type,
        extractedText=normalized_text,
        characterCount=len(normalized_text),
        detectedEmails=_detect_emails(normalized_text),
        detectedPhones=_detect_phones(normalized_text),
        detectedSections=_detect_sections(normalized_text),
    )


def _extract_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="ignore")


def _extract_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not extract text from PDF") from exc


def _extract_docx(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_text = []
        for table in document.tables:
            for row in table.rows:
                table_text.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs + table_text)
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not extract text from DOCX") from exc


def _normalize_text(text: str) -> str:
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def _detect_emails(text: str) -> list[str]:
    return sorted(set(re.findall(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)))


def _detect_phones(text: str) -> list[str]:
    matches = re.findall(r"(?:\+?\d[\d\s().-]{8,}\d)", text)
    return sorted(set(match.strip() for match in matches))


def _detect_sections(text: str) -> list[str]:
    known_sections = [
        "experience",
        "projects",
        "certifications",
        "achievements",
        "skills",
        "education",
        "summary",
        "profile",
    ]
    lowered = text.lower()
    return [section for section in known_sections if section in lowered]
